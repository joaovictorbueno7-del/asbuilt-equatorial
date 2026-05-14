"""Gerador de Relatório de Entrega de Obra — formato Equatorial.

Produz DOCX (python-docx) e PDF (reportlab) no layout oficial:
  - Cabeçalho em TODAS as páginas: logo esq | título centralizado | parceira dir | linha azul
  - Página 1: Nota / Município / Parceira em negrito + título "Postes, Estruturas e Redes"
  - Grade de fotos: 2 por linha, legenda em negrito 10pt centralizada
"""
from __future__ import annotations

import io
import os
import re
import zipfile
from pathlib import Path

from loguru import logger

# ── python-docx ───────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── reportlab ─────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
    Table, TableStyle, HRFlowable,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# ─────────────────────────────────────────────────────────────────────────────
# Constantes
# ─────────────────────────────────────────────────────────────────────────────

# Azul Equatorial
_BLUE = RGBColor(0x00, 0x5F, 0xAF)
_BLUE_HEX = "005FAF"

_MAX_IMG_BYTES = 8 * 1024 * 1024  # 8 MB


def _reports_dir() -> Path:
    """Diretório de saída — avaliado em tempo de execução para respeitar env vars."""
    d = Path(os.environ.get("STORAGE_LOCAL_PATH", "storage")) / "reports"
    d.mkdir(parents=True, exist_ok=True)
    return d


# ─────────────────────────────────────────────────────────────────────────────
# Helpers de imagem
# ─────────────────────────────────────────────────────────────────────────────

def _load_image_bytes(src: str | None, kmz_path: str | None) -> bytes | None:
    """Carrega bytes de uma imagem: caminho absoluto ou dentro do KMZ."""
    if not src:
        return None
    p = Path(src)
    if p.is_file():
        data = p.read_bytes()
        return data if len(data) <= _MAX_IMG_BYTES else None
    if kmz_path and Path(kmz_path).is_file():
        safe = src.replace("..", "").lstrip("/").lstrip("\\")
        try:
            with zipfile.ZipFile(kmz_path, "r") as zf:
                names = zf.namelist()
                target = next(
                    (n for n in names if n == safe or Path(n).name == Path(safe).name),
                    None,
                )
                if target:
                    return zf.read(target)
        except Exception:
            pass
    return None


def _resize_image(data: bytes, max_w: int = 1200, max_h: int = 900) -> bytes:
    """Redimensiona imagem se necessária (mantém proporção)."""
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(data))
        img.thumbnail((max_w, max_h), Image.LANCZOS)
        buf = io.BytesIO()
        fmt = img.format or "JPEG"
        if fmt not in ("JPEG", "PNG"):
            fmt = "JPEG"
            img = img.convert("RGB")
        img.save(buf, format=fmt, quality=85)
        return buf.getvalue()
    except Exception:
        return data


# ─────────────────────────────────────────────────────────────────────────────
# Extração de metadados do KML
# ─────────────────────────────────────────────────────────────────────────────

def _extract_kml_metadata(kmz_path: str | None) -> dict:
    """Extrai metadados do doc.kml dentro do KMZ.

    Retorna dict com: nome_projeto, coordinates (list), placemark_names
    """
    result: dict = {}
    if not kmz_path or not Path(kmz_path).is_file():
        return result
    try:
        with zipfile.ZipFile(kmz_path, "r") as zf:
            kml_names = [n for n in zf.namelist() if n.endswith(".kml")]
            if not kml_names:
                return result
            kml_text = zf.read(kml_names[0]).decode("utf-8", errors="replace")

        # Nome do projeto / Folder name
        m = re.search(r"<Folder[^>]*>.*?<name>(.*?)</name>", kml_text, re.DOTALL)
        if m:
            result["nome_projeto"] = m.group(1).strip()
        else:
            m2 = re.search(r"<Document[^>]*>.*?<name>(.*?)</name>", kml_text, re.DOTALL)
            if m2:
                result["nome_projeto"] = m2.group(1).strip()

        # Coordenadas de todos os placemarks (lon,lat)
        coords = re.findall(r"<coordinates>([\-\d\.]+),([\-\d\.]+)", kml_text)
        result["coordinates"] = [(float(c[0]), float(c[1])) for c in coords]

        # Nomes dos placemarks
        result["placemark_names"] = re.findall(r"<Placemark[^>]*>.*?<name>(.*?)</name>",
                                               kml_text, re.DOTALL)

    except Exception as e:
        logger.warning(f"[report_builder] erro ao ler KML: {e}")
    return result


# ─────────────────────────────────────────────────────────────────────────────
# Agrupamento de estruturas
# ─────────────────────────────────────────────────────────────────────────────

def _group_structures(structures: list[dict]) -> list[tuple[str, list[dict]]]:
    """Agrupa estruturas por prefixo do nome do placemark (tipo do poste)."""
    groups: dict[str, list[dict]] = {}
    for s in structures:
        pname = s.get("placemark_name", "")
        parts = pname.split(".")
        group_key = parts[0].strip() if parts else pname
        if not group_key:
            group_key = s.get("structure_type", "Outras Estruturas")
        groups.setdefault(group_key, []).append(s)
    return list(groups.items())


def _extract_poste_type(name: str) -> str:
    """Extrai tipo do poste do nome do placemark (ex: 'PDT 9/300')."""
    m = re.search(r"\b(P?[A-Z]{2,3})\s+(\d+/\d+)", name, re.IGNORECASE)
    if m:
        return f"Poste {m.group(1).upper()} {m.group(2)}"
    m2 = re.search(r"\b([A-Z]{2,3})\s+(\d+)", name, re.IGNORECASE)
    if m2:
        return f"Poste {m2.group(1).upper()} {m2.group(2)}"
    return "Poste"


# ─────────────────────────────────────────────────────────────────────────────
# DOCX — helpers XML
# ─────────────────────────────────────────────────────────────────────────────

def _remove_cell_borders(cell) -> None:
    """Remove todas as bordas de uma célula DOCX via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove tcBorders existente se houver
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tag.set(qn("w:sz"), "0")
        tag.set(qn("w:color"), "auto")
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def _set_cell_width(cell, width_cm: float) -> None:
    """Define largura de uma célula DOCX em centímetros."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))  # 1cm = 567 twips
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_table_layout_fixed(tbl) -> None:
    """Força layout fixo na tabela DOCX."""
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)


def _add_hr_paragraph(container, color_hex: str = _BLUE_HEX,
                      space_before: float = 4.0, space_after: float = 4.0) -> None:
    """Adiciona parágrafo com linha separadora inferior (borda w:bottom)."""
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), color_hex)
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_meta_line(doc: Document, label: str, value: str, size_pt: float = 11.0) -> None:
    """Adiciona linha 'Label: valor' com label em negrito e valor normal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r_lbl = p.add_run(f"{label}: ")
    r_lbl.bold = True
    r_lbl.font.size = Pt(size_pt)
    r_val = p.add_run(value or "—")
    r_val.font.size = Pt(size_pt)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX — Cabeçalho
# ─────────────────────────────────────────────────────────────────────────────

def _add_header(doc: Document, tipo: str, parceira: str) -> None:
    """Monta o cabeçalho Equatorial em todas as páginas.

    Estrutura:
      [tabela sem bordas: logo esq | título central | parceira dir]
      [parágrafo com borda inferior azul = linha separadora]

    A tabela é inserida ANTES do parágrafo padrão do cabeçalho usando
    lxml.addprevious(), garantindo a ordem correta independente da API
    do python-docx.
    """
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Limpamos o parágrafo padrão (mas o mantemos — será o separador)
    default_para = header.paragraphs[0]
    default_para.clear()

    # ── Tabela 1×3 ────────────────────────────────────────────────────────────
    # Largura total = 17 cm (A4 − margens 2+2)
    COL_LOGO = 3.5   # cm
    COL_TITLE = 10.0  # cm
    COL_PARTNER = 3.5  # cm

    htable = header.add_table(rows=1, cols=3, width=Cm(COL_LOGO + COL_TITLE + COL_PARTNER))
    htable.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_layout_fixed(htable)

    # Coloca tabela ANTES do parágrafo separador (reordenação via lxml)
    default_para._p.addprevious(htable._tbl)

    # Larguras das colunas
    _set_cell_width(htable.cell(0, 0), COL_LOGO)
    _set_cell_width(htable.cell(0, 1), COL_TITLE)
    _set_cell_width(htable.cell(0, 2), COL_PARTNER)

    # Remover bordas
    for col in range(3):
        _remove_cell_borders(htable.cell(0, col))

    # ── Coluna 0: Logo / "equatorial" ─────────────────────────────────────────
    c0 = htable.cell(0, 0)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after = Pt(2)
    r0 = p0.add_run("equatorial")
    r0.bold = True
    r0.font.size = Pt(12)
    r0.font.color.rgb = _BLUE

    # ── Coluna 1: Título (3 linhas) ───────────────────────────────────────────
    c1 = htable.cell(0, 1)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Linha 1: "Relatório de Entrega de Obra" — negrito
    p1a = c1.paragraphs[0]
    p1a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1a.paragraph_format.space_before = Pt(2)
    p1a.paragraph_format.space_after = Pt(0)
    r1a = p1a.add_run("Relatório de Entrega de Obra")
    r1a.bold = True
    r1a.font.size = Pt(9)

    # Linha 2: tipo da obra
    p1b = c1.add_paragraph()
    p1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1b.paragraph_format.space_before = Pt(0)
    p1b.paragraph_format.space_after = Pt(0)
    r1b = p1b.add_run(tipo or "")
    r1b.font.size = Pt(8)

    # Linha 3: regional
    p1c = c1.add_paragraph()
    p1c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1c.paragraph_format.space_before = Pt(0)
    p1c.paragraph_format.space_after = Pt(2)
    r1c = p1c.add_run("Superintendência Centro – Regional Metropolitana")
    r1c.font.size = Pt(7)
    r1c.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Coluna 2: Parceira ────────────────────────────────────────────────────
    c2 = htable.cell(0, 2)
    c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(parceira[:30] if parceira else "")
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── Linha separadora (no parágrafo após a tabela) ─────────────────────────
    default_para.paragraph_format.space_before = Pt(0)
    default_para.paragraph_format.space_after = Pt(0)
    pPr = default_para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")   # 1.5pt
    bottom.set(qn("w:color"), _BLUE_HEX)
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX — build principal
# ─────────────────────────────────────────────────────────────────────────────

def build_docx(
    run_id: str,
    metadata: dict,
    structures: list[dict],
    kmz_path: str | None = None,
) -> bytes:
    """Constrói o DOCX no formato oficial Equatorial."""
    nota = metadata.get("nota") or "—"
    municipio = metadata.get("municipio") or "—"
    parceira = metadata.get("parceira") or "—"
    tipo = metadata.get("tipo") or "Postes, Estruturas e Redes"

    doc = Document()

    # ── Configuração de página (A4) ───────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(3.2)   # espaço para o cabeçalho
    section.bottom_margin = Cm(2.0)

    # ── Cabeçalho ─────────────────────────────────────────────────────────────
    _add_header(doc, tipo, parceira)

    # ── Metadados (Página 1) ──────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(4)  # espaçamento inicial
    _add_meta_line(doc, "Nota", nota)
    _add_meta_line(doc, "Município", municipio)
    _add_meta_line(doc, "Parceira Construção", parceira)

    # ── Título "Postes, Estruturas e Redes" com separadores ───────────────────
    _add_hr_paragraph(doc, space_before=8, space_after=6)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    rt = p_title.add_run("Postes, Estruturas e Redes")
    rt.bold = True
    rt.font.size = Pt(13)

    _add_hr_paragraph(doc, space_before=6, space_after=10)

    # ── Grade de fotos ────────────────────────────────────────────────────────
    # Largura útil: 17cm; cada foto ocupa ~8cm
    PHOTO_W_INCHES = Inches(3.0)   # ~7.6cm — cabe bem em 2 colunas
    COL_W_CM = 8.5                  # cada coluna = metade de 17cm

    groups = _group_structures(structures)
    for group_title, group_structs in groups:
        # Título do grupo (negrito, centralizado)
        gp = doc.add_paragraph()
        gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gp.paragraph_format.space_before = Pt(6)
        gp.paragraph_format.space_after = Pt(4)
        gr = gp.add_run(group_title)
        gr.bold = True
        gr.font.size = Pt(11)

        # Pares de fotos
        for i in range(0, len(group_structs), 2):
            pair = group_structs[i: i + 2]

            # Sempre 2 colunas para manter alinhamento
            tbl = doc.add_table(rows=2, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_layout_fixed(tbl)

            for col_idx in range(2):
                # Definir larguras
                _set_cell_width(tbl.cell(0, col_idx), COL_W_CM)
                _set_cell_width(tbl.cell(1, col_idx), COL_W_CM)
                _remove_cell_borders(tbl.cell(0, col_idx))
                _remove_cell_borders(tbl.cell(1, col_idx))

                if col_idx >= len(pair):
                    # Célula vazia (último par com 1 foto)
                    continue

                struct = pair[col_idx]
                img_src = struct.get("image_src") or struct.get("photo_src")
                caption = struct.get("caption") or struct.get("placemark_name") or "—"

                # Linha 0: foto
                cell_img = tbl.cell(0, col_idx)
                cell_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cp = cell_img.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_before = Pt(2)
                cp.paragraph_format.space_after = Pt(2)

                img_bytes = _load_image_bytes(img_src, kmz_path)
                if img_bytes:
                    try:
                        img_bytes = _resize_image(img_bytes)
                        cp.add_run().add_picture(io.BytesIO(img_bytes), width=PHOTO_W_INCHES)
                    except Exception as e:
                        logger.warning(f"[report_builder] foto falhou ({img_src}): {e}")
                        cp.add_run("[foto indisponível]").font.size = Pt(9)
                else:
                    cp.add_run("[foto indisponível]").font.size = Pt(9)

                # Linha 1: legenda — negrito, centralizado, 10pt
                cell_cap = tbl.cell(1, col_idx)
                lp = cell_cap.paragraphs[0]
                lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lp.paragraph_format.space_before = Pt(2)
                lp.paragraph_format.space_after = Pt(4)
                lr = lp.add_run(caption)
                lr.bold = True
                lr.font.size = Pt(10)

            # Espaço entre pares
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(6)

    # ── Salva ─────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# PDF Builder (reportlab)
# ─────────────────────────────────────────────────────────────────────────────

def build_pdf(
    run_id: str,
    metadata: dict,
    structures: list[dict],
    kmz_path: str | None = None,
) -> bytes:
    """Constrói PDF no formato Equatorial usando reportlab."""
    nota = metadata.get("nota") or "—"
    municipio = metadata.get("municipio") or "—"
    parceira = metadata.get("parceira") or "—"
    tipo = metadata.get("tipo") or "Postes, Estruturas e Redes"

    buf = io.BytesIO()
    PAGE_W, PAGE_H = A4
    MARGIN = 2 * cm
    HDR_H = 2.0 * cm  # altura do cabeçalho

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=MARGIN,
        rightMargin=MARGIN,
        topMargin=HDR_H + 1.2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=11, leading=14, spaceAfter=3)
    body_bold = ParagraphStyle("BodyBold", parent=styles["Normal"], fontSize=11, leading=14,
                                fontName="Helvetica-Bold", spaceAfter=3)
    h1 = ParagraphStyle("H1", parent=styles["Normal"], fontSize=13, leading=16,
                         alignment=TA_CENTER, fontName="Helvetica-Bold", spaceAfter=4, spaceBefore=0)
    h2 = ParagraphStyle("H2", parent=styles["Normal"], fontSize=11, leading=14,
                         alignment=TA_CENTER, fontName="Helvetica-Bold", spaceAfter=4, spaceBefore=8)
    caption_style = ParagraphStyle("Caption", parent=styles["Normal"], fontSize=10,
                                    leading=13, alignment=TA_CENTER, fontName="Helvetica-Bold")

    # Callbacks para cabeçalho
    _hdr_meta = {
        "equatorial": "equatorial",
        "title_lines": ["Relatório de Entrega de Obra", tipo,
                        "Superintendência Centro – Regional Metropolitana"],
        "partner": parceira[:30] if parceira else "",
        "page_w": PAGE_W,
        "page_h": PAGE_H,
        "hdr_h": HDR_H,
        "margin": MARGIN,
    }

    def _on_page(canvas, doc_obj):
        _draw_pdf_header(canvas, _hdr_meta)

    # ── Conteúdo ──────────────────────────────────────────────────────────────
    content = []

    # Metadados com label em negrito
    def meta_row(label, value):
        return Paragraph(f"<b>{label}:</b> {value or '—'}", body)

    content.append(meta_row("Nota", nota))
    content.append(meta_row("Município", municipio))
    content.append(meta_row("Parceira Construção", parceira))
    content.append(Spacer(1, 8))
    content.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor(f"#{_BLUE_HEX}")))
    content.append(Spacer(1, 4))
    content.append(Paragraph("Postes, Estruturas e Redes", h1))
    content.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor(f"#{_BLUE_HEX}")))
    content.append(Spacer(1, 12))

    # Grade de fotos
    photo_w = (PAGE_W - 2 * MARGIN - 0.5 * cm) * 0.47
    photo_h = photo_w * 0.75

    groups = _group_structures(structures)
    for group_title, group_structs in groups:
        content.append(Paragraph(group_title, h2))

        for i in range(0, len(group_structs), 2):
            pair = group_structs[i: i + 2]
            row_imgs: list = []
            row_caps: list = []

            for j in range(2):
                if j < len(pair):
                    struct = pair[j]
                    img_src = struct.get("image_src") or struct.get("photo_src")
                    cap_text = struct.get("caption") or struct.get("placemark_name") or "—"
                    img_bytes = _load_image_bytes(img_src, kmz_path)
                    if img_bytes:
                        try:
                            img_bytes = _resize_image(img_bytes, 1200, 900)
                            row_imgs.append(RLImage(io.BytesIO(img_bytes),
                                                    width=photo_w, height=photo_h))
                        except Exception as e:
                            logger.warning(f"[report_builder/pdf] foto falhou: {e}")
                            row_imgs.append(Paragraph("[foto indisponível]", caption_style))
                    else:
                        row_imgs.append(Paragraph("[foto indisponível]", caption_style))
                    row_caps.append(Paragraph(cap_text, caption_style))
                else:
                    # Célula vazia
                    row_imgs.append(Spacer(photo_w, photo_h))
                    row_caps.append(Paragraph("", caption_style))

            col_w = (PAGE_W - 2 * MARGIN) / 2
            tbl = Table([row_imgs, row_caps], colWidths=[col_w, col_w])
            tbl.setStyle(TableStyle([
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, 0), "MIDDLE"),
                ("VALIGN", (0, 1), (-1, 1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.lightgrey),
            ]))
            content.append(tbl)
            content.append(Spacer(1, 8))

    doc.build(content, onFirstPage=_on_page, onLaterPages=_on_page)
    return buf.getvalue()


def _draw_pdf_header(canvas, meta: dict) -> None:
    """Desenha cabeçalho fixo em cada página PDF."""
    from reportlab.lib.units import cm as rl_cm

    canvas.saveState()
    page_w = meta["page_w"]
    page_h = meta["page_h"]
    hdr_h = meta["hdr_h"]
    margin = meta["margin"]
    usable_w = page_w - 2 * margin

    top_y = page_h - 0.5 * rl_cm
    bot_y = top_y - hdr_h

    # Fundo levemente cinza
    canvas.setFillColorRGB(0.97, 0.97, 0.97)
    canvas.rect(margin, bot_y, usable_w, hdr_h, fill=1, stroke=0)

    mid_y = (top_y + bot_y) / 2

    # ── Logo esquerda ─────────────────────────────────────────────────────────
    canvas.setFont("Helvetica-Bold", 12)
    canvas.setFillColorRGB(0.0, 0.373, 0.686)
    canvas.drawString(margin + 6, mid_y - 5, meta["equatorial"])

    # ── Título centralizado ───────────────────────────────────────────────────
    lines = meta["title_lines"]
    cx = page_w / 2
    if len(lines) >= 1:
        canvas.setFont("Helvetica-Bold", 9)
        canvas.setFillColorRGB(0.1, 0.1, 0.1)
        canvas.drawCentredString(cx, mid_y + 5, lines[0])
    if len(lines) >= 2:
        canvas.setFont("Helvetica", 8)
        canvas.setFillColorRGB(0.2, 0.2, 0.2)
        canvas.drawCentredString(cx, mid_y - 5, lines[1])
    if len(lines) >= 3:
        canvas.setFont("Helvetica", 7)
        canvas.setFillColorRGB(0.45, 0.45, 0.45)
        canvas.drawCentredString(cx, mid_y - 14, lines[2])

    # ── Parceira direita ──────────────────────────────────────────────────────
    canvas.setFont("Helvetica", 8)
    canvas.setFillColorRGB(0.4, 0.4, 0.4)
    canvas.drawRightString(page_w - margin - 6, mid_y - 5, meta["partner"])

    # ── Linha separadora azul ─────────────────────────────────────────────────
    canvas.setStrokeColorRGB(0.0, 0.373, 0.686)
    canvas.setLineWidth(1.5)
    canvas.line(margin, bot_y - 1, page_w - margin, bot_y - 1)

    canvas.restoreState()


# ─────────────────────────────────────────────────────────────────────────────
# Extração de estruturas do AgentRun output
# ─────────────────────────────────────────────────────────────────────────────

def extract_structures_from_run(output_payload: dict, kmz_path: str | None = None) -> list[dict]:
    """Converte output_payload do AgentRun em lista de estruturas para o relatório."""
    structures = []
    placemarks = output_payload.get("placemarks") or output_payload.get("structures") or []

    for pm in placemarks:
        name = pm.get("name") or pm.get("placemark_name") or "—"
        photos = pm.get("photos") or pm.get("images") or []

        # Códigos de estrutura confirmados ou declarados
        struct_codes: list[str] = []
        confirmed = pm.get("estruturas_confirmadas") or []
        declared = pm.get("declared_codes") or []
        if confirmed:
            struct_codes = confirmed[:3]
        elif pm.get("structure_type"):
            struct_codes = [pm["structure_type"]]
        elif declared:
            struct_codes = declared[:3]

        poste_type = _extract_poste_type(name)
        codes_str = " ".join(struct_codes) if struct_codes else ""
        caption = f"{poste_type} - {codes_str}".strip(" -") if codes_str else poste_type

        if photos:
            for photo in photos:
                photo_src = photo if isinstance(photo, str) else (
                    photo.get("src") or photo.get("path") or photo.get("file_rel_path", "")
                )
                structures.append({
                    "placemark_name": name,
                    "image_src": photo_src,
                    "caption": caption,
                })
        else:
            structures.append({
                "placemark_name": name,
                "image_src": None,
                "caption": caption,
            })

    return structures


# ─────────────────────────────────────────────────────────────────────────────
# Mapeamento tipo código → label legível para o cabeçalho
# ─────────────────────────────────────────────────────────────────────────────

_TIPO_LABELS: dict[str, str] = {
    "as_built":            "AS BUILT",
    "obras":               "Construção",
    "construcao":          "Construção",
    "manutencao":          "Manutenção",
    "manutencao_corretiva":"Manutenção Corretiva",
    "inspecao":            "Inspeção",
    "entrega":             "Entrega de Obra",
}


def _tipo_label(raw: str | None) -> str:
    """Converte código do tipo (ex: 'as_built') para label de exibição."""
    if not raw:
        return "AS BUILT"
    raw_clean = raw.strip().lower().replace(" ", "_").replace("-", "_")
    # Busca exata
    if raw_clean in _TIPO_LABELS:
        return _TIPO_LABELS[raw_clean]
    # Se já é um texto legível (ex: "Construção"), retorna como está
    if any(c.isalpha() and c == c.upper() for c in raw):
        return raw
    return raw.upper()


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def generate_report(
    run_id: str,
    output_payload: dict,
    kmz_path: str | None = None,
    fmt: str = "docx",
    filename_override: str | None = None,
) -> tuple[bytes, str]:
    """Gera relatório a partir do output de um AgentRun.

    Args:
        run_id: ID do run (usado no nome do arquivo)
        output_payload: dict com placemarks/structures + metadados
        kmz_path: caminho do KMZ para extração de fotos
        fmt: "docx" ou "pdf"
        filename_override: nome do arquivo de saída (sem extensão)

    Returns:
        (bytes_do_arquivo, nome_do_arquivo)
    """
    # Tenta extrair metadados do KML se disponíveis
    kml_meta = _extract_kml_metadata(kmz_path)

    # Nota: usa work_name > nome_projeto do KML > run_id[:8]
    nota = (
        output_payload.get("nota")
        or output_payload.get("work_name")
        or kml_meta.get("nome_projeto")
        or run_id[:8]
    )

    # tipo: código interno (as_built, obras, manutencao) → label legível para cabeçalho
    tipo_raw = output_payload.get("tipo") or "as_built"
    tipo_label = _tipo_label(tipo_raw)

    metadata = {
        "nota": nota,
        "municipio": output_payload.get("municipio") or "—",
        "parceira": (
            output_payload.get("parceira")
            or output_payload.get("concessionaria")
            or "—"
        ),
        "tipo": tipo_label,  # somente para o cabeçalho; corpo usa "Postes, Estruturas e Redes"
    }

    structures = extract_structures_from_run(output_payload, kmz_path)
    logger.info(f"[report_builder] run={run_id} fmt={fmt} estruturas={len(structures)}")

    stem = filename_override or f"relatorio_{run_id[:8]}"
    # Garante que o stem não tem extensão
    stem = stem.removesuffix(".docx").removesuffix(".pdf")

    if fmt == "pdf":
        data = build_pdf(run_id, metadata, structures, kmz_path)
        filename = f"{stem}.pdf"
    else:
        data = build_docx(run_id, metadata, structures, kmz_path)
        filename = f"{stem}.docx"

    out_path = _reports_dir() / filename
    out_path.write_bytes(data)
    logger.info(f"[report_builder] salvo em {out_path} ({len(data):,} bytes)")

    return data, filename
